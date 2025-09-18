#!/usr/bin/env python3
"""
Test script to create documents for plagiarism detection testing
"""

from docx import Document
import os

def create_original_document():
    """Create a document with original content"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('The Impact of Artificial Intelligence on Modern Education', 0)
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Artificial Intelligence (AI) has revolutionized numerous industries, and education is no exception. '
        'The integration of AI technologies in educational settings has brought about significant changes '
        'in how students learn, how teachers teach, and how educational institutions operate. This paper '
        'explores the multifaceted impact of AI on modern education, examining both the opportunities '
        'and challenges it presents.'
    )
    
    # Add main content
    doc.add_heading('AI-Powered Learning Systems', level=1)
    doc.add_paragraph(
        'One of the most significant contributions of AI to education is the development of intelligent '
        'learning systems. These systems can adapt to individual student needs, providing personalized '
        'learning experiences that were previously impossible to achieve at scale. Machine learning '
        'algorithms analyze student performance data to identify learning patterns and adjust content '
        'delivery accordingly.'
    )
    
    doc.add_paragraph(
        'For instance, adaptive learning platforms use AI to create customized learning paths for each '
        'student. These systems can identify when a student is struggling with a particular concept '
        'and provide additional resources or alternative explanations. This personalized approach '
        'has been shown to improve learning outcomes and student engagement significantly.'
    )
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The integration of AI in education represents a paradigm shift that offers tremendous potential '
        'for improving learning outcomes. However, it also requires careful consideration of ethical '
        'implications, data privacy concerns, and the need for proper teacher training. As AI continues '
        'to evolve, educational institutions must strike a balance between embracing innovation and '
        'maintaining the human elements that are essential to effective education.'
    )
    
    # Save the document
    doc.save('original_ai_education.docx')
    print("✅ Created original_ai_education.docx")
    return 'original_ai_education.docx'

def create_plagiarized_document():
    """Create a document with plagiarized content (copied from the original)"""
    doc = Document()
    
    # Add title (slightly modified)
    title = doc.add_heading('How AI is Transforming Education Today', 0)
    
    # Add introduction (copied with minor changes)
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Artificial Intelligence (AI) has revolutionized numerous industries, and education is no exception. '
        'The integration of AI technologies in educational settings has brought about significant changes '
        'in how students learn, how teachers teach, and how educational institutions operate. This paper '
        'explores the multifaceted impact of AI on modern education, examining both the opportunities '
        'and challenges it presents.'
    )
    
    # Add main content (copied with minor modifications)
    doc.add_heading('AI-Powered Learning Systems', level=1)
    doc.add_paragraph(
        'One of the most significant contributions of AI to education is the development of intelligent '
        'learning systems. These systems can adapt to individual student needs, providing personalized '
        'learning experiences that were previously impossible to achieve at scale. Machine learning '
        'algorithms analyze student performance data to identify learning patterns and adjust content '
        'delivery accordingly.'
    )
    
    doc.add_paragraph(
        'For instance, adaptive learning platforms use AI to create customized learning paths for each '
        'student. These systems can identify when a student is struggling with a particular concept '
        'and provide additional resources or alternative explanations. This personalized approach '
        'has been shown to improve learning outcomes and student engagement significantly.'
    )
    
    # Add conclusion (copied with minor changes)
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The integration of AI in education represents a paradigm shift that offers tremendous potential '
        'for improving learning outcomes. However, it also requires careful consideration of ethical '
        'implications, data privacy concerns, and the need for proper teacher training. As AI continues '
        'to evolve, educational institutions must strike a balance between embracing innovation and '
        'maintaining the human elements that are essential to effective education.'
    )
    
    # Save the document
    doc.save('plagiarized_ai_education.docx')
    print("✅ Created plagiarized_ai_education.docx")
    return 'plagiarized_ai_education.docx'

def create_completely_different_document():
    """Create a document with completely different content"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('The Future of Renewable Energy Technologies', 0)
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Renewable energy technologies have emerged as the cornerstone of sustainable development '
        'in the 21st century. As the world grapples with climate change and the depletion of fossil '
        'fuels, renewable energy sources such as solar, wind, and hydroelectric power have become '
        'increasingly important. This paper examines the current state of renewable energy technologies '
        'and their potential to transform the global energy landscape.'
    )
    
    # Add main content
    doc.add_heading('Solar Energy Innovations', level=1)
    doc.add_paragraph(
        'Solar energy has experienced remarkable technological advancements in recent years. '
        'Photovoltaic cells have become more efficient and cost-effective, making solar power '
        'accessible to a broader range of consumers. The development of thin-film solar panels '
        'and concentrated solar power systems has further expanded the applications of solar energy.'
    )
    
    doc.add_paragraph(
        'Wind energy technology has also seen significant improvements. Modern wind turbines '
        'are more efficient and can generate electricity even in low-wind conditions. Offshore '
        'wind farms have emerged as a promising solution for countries with limited land resources, '
        'providing access to stronger and more consistent wind patterns.'
    )
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The future of renewable energy technologies looks promising, with continuous innovations '
        'driving down costs and improving efficiency. However, challenges such as energy storage, '
        'grid integration, and policy support remain critical factors in determining the success '
        'of renewable energy adoption. Continued investment in research and development will be '
        'essential to overcome these challenges and achieve a sustainable energy future.'
    )
    
    # Save the document
    doc.save('different_renewable_energy.docx')
    print("✅ Created different_renewable_energy.docx")
    return 'different_renewable_energy.docx'

def main():
    """Create all test documents"""
    print("🔧 Creating test documents for plagiarism detection...")
    print()
    
    # Create original document
    original_file = create_original_document()
    
    # Create plagiarized document
    plagiarized_file = create_plagiarized_document()
    
    # Create completely different document
    different_file = create_completely_different_document()
    
    print()
    print("📋 Test Documents Created:")
    print(f"   📄 {original_file} - Original content about AI in education")
    print(f"   📄 {plagiarized_file} - Plagiarized content (should show high similarity)")
    print(f"   📄 {different_file} - Different content about renewable energy")
    print()
    print("🎯 Now you can test plagiarism detection:")
    print("   1. Upload the original document first")
    print("   2. Then upload the plagiarized document")
    print("   3. Compare the similarity scores!")
    print("   4. Upload the different document to see low similarity")

if __name__ == "__main__":
    main()

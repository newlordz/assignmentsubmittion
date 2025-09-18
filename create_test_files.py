#!/usr/bin/env python3
"""
Create simple test files for plagiarism detection
"""

from docx import Document
import os

def create_original_assignment():
    """Create an original assignment document"""
    doc = Document()
    
    # Add title
    doc.add_heading('The Benefits of Online Learning', 0)
    
    # Add student info
    doc.add_paragraph('Student: John Smith')
    doc.add_paragraph('Course: Introduction to Education')
    doc.add_paragraph('Date: September 2024')
    doc.add_paragraph('')
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Online learning has become increasingly popular in recent years, especially after the global pandemic. '
        'This educational approach offers numerous advantages for both students and educators. In this essay, '
        'I will explore the key benefits of online learning and how it has transformed the educational landscape.'
    )
    
    # Add main content
    doc.add_heading('Flexibility and Convenience', level=1)
    doc.add_paragraph(
        'One of the most significant advantages of online learning is the flexibility it provides. Students '
        'can access course materials and complete assignments at their own pace and schedule. This is particularly '
        'beneficial for working professionals, parents, and individuals with busy lifestyles who cannot attend '
        'traditional classroom sessions.'
    )
    
    doc.add_paragraph(
        'Additionally, online learning eliminates the need for commuting to campus, saving both time and money. '
        'Students can study from the comfort of their homes or any location with internet access, making education '
        'more accessible to people in remote areas or those with mobility challenges.'
    )
    
    # Add more content
    doc.add_heading('Cost-Effectiveness', level=1)
    doc.add_paragraph(
        'Online education is often more cost-effective than traditional classroom learning. Students can save '
        'money on transportation, accommodation, and other campus-related expenses. Many online courses also '
        'offer digital textbooks and resources, reducing the cost of educational materials.'
    )
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'In conclusion, online learning offers numerous benefits including flexibility, convenience, and '
        'cost-effectiveness. While it may not be suitable for every student or subject, it has proven to be '
        'a valuable alternative to traditional education. As technology continues to advance, online learning '
        'will likely become even more sophisticated and widely adopted.'
    )
    
    # Save the document
    doc.save('original_online_learning.docx')
    print("✅ Created original_online_learning.docx")
    return 'original_online_learning.docx'

def create_plagiarized_assignment():
    """Create a plagiarized assignment document (copied from the original)"""
    doc = Document()
    
    # Add title (slightly modified)
    doc.add_heading('Advantages of Digital Education', 0)
    
    # Add student info (different student)
    doc.add_paragraph('Student: Sarah Johnson')
    doc.add_paragraph('Course: Educational Technology')
    doc.add_paragraph('Date: September 2024')
    doc.add_paragraph('')
    
    # Add introduction (copied with minor changes)
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Online learning has become increasingly popular in recent years, especially after the global pandemic. '
        'This educational approach offers numerous advantages for both students and educators. In this essay, '
        'I will explore the key benefits of online learning and how it has transformed the educational landscape.'
    )
    
    # Add main content (copied with minor modifications)
    doc.add_heading('Flexibility and Convenience', level=1)
    doc.add_paragraph(
        'One of the most significant advantages of online learning is the flexibility it provides. Students '
        'can access course materials and complete assignments at their own pace and schedule. This is particularly '
        'beneficial for working professionals, parents, and individuals with busy lifestyles who cannot attend '
        'traditional classroom sessions.'
    )
    
    doc.add_paragraph(
        'Additionally, online learning eliminates the need for commuting to campus, saving both time and money. '
        'Students can study from the comfort of their homes or any location with internet access, making education '
        'more accessible to people in remote areas or those with mobility challenges.'
    )
    
    # Add more content (copied)
    doc.add_heading('Cost-Effectiveness', level=1)
    doc.add_paragraph(
        'Online education is often more cost-effective than traditional classroom learning. Students can save '
        'money on transportation, accommodation, and other campus-related expenses. Many online courses also '
        'offer digital textbooks and resources, reducing the cost of educational materials.'
    )
    
    # Add conclusion (copied with minor changes)
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'In conclusion, online learning offers numerous benefits including flexibility, convenience, and '
        'cost-effectiveness. While it may not be suitable for every student or subject, it has proven to be '
        'a valuable alternative to traditional education. As technology continues to advance, online learning '
        'will likely become even more sophisticated and widely adopted.'
    )
    
    # Save the document
    doc.save('plagiarized_online_learning.docx')
    print("✅ Created plagiarized_online_learning.docx")
    return 'plagiarized_online_learning.docx'

def create_completely_different_assignment():
    """Create a completely different assignment"""
    doc = Document()
    
    # Add title
    doc.add_heading('The Impact of Climate Change on Agriculture', 0)
    
    # Add student info
    doc.add_paragraph('Student: Michael Brown')
    doc.add_paragraph('Course: Environmental Science')
    doc.add_paragraph('Date: September 2024')
    doc.add_paragraph('')
    
    # Add introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Climate change is one of the most pressing challenges of our time, with far-reaching implications '
        'for various sectors including agriculture. Rising temperatures, changing precipitation patterns, '
        'and extreme weather events are significantly affecting crop yields and farming practices worldwide. '
        'This essay examines the impact of climate change on agriculture and potential adaptation strategies.'
    )
    
    # Add main content
    doc.add_heading('Temperature Effects on Crop Growth', level=1)
    doc.add_paragraph(
        'Rising global temperatures have both positive and negative effects on agricultural productivity. '
        'While some crops may benefit from warmer conditions in certain regions, many staple crops such as '
        'wheat, rice, and corn are sensitive to temperature changes. Studies show that for every degree '
        'Celsius increase in temperature, crop yields can decrease by 3-5% for major cereal crops.'
    )
    
    doc.add_paragraph(
        'Extreme heat events can cause heat stress in plants, leading to reduced photosynthesis and '
        'poor crop quality. Additionally, higher temperatures can accelerate the life cycles of pests '
        'and diseases, increasing the need for pest management strategies.'
    )
    
    # Add more content
    doc.add_heading('Water Availability and Irrigation', level=1)
    doc.add_paragraph(
        'Climate change is altering precipitation patterns, leading to more frequent droughts in some '
        'regions and increased flooding in others. These changes directly impact water availability for '
        'agriculture. Farmers must adapt by implementing more efficient irrigation systems and water '
        'conservation techniques.'
    )
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'Climate change poses significant challenges to global agriculture, requiring immediate action '
        'and long-term adaptation strategies. Farmers, researchers, and policymakers must work together '
        'to develop resilient agricultural systems that can withstand changing climatic conditions. '
        'Investment in research, technology, and sustainable farming practices is essential for '
        'ensuring food security in a changing climate.'
    )
    
    # Save the document
    doc.save('different_climate_agriculture.docx')
    print("✅ Created different_climate_agriculture.docx")
    return 'different_climate_agriculture.docx'

def main():
    """Create all test files"""
    print("🔧 Creating Test Files for Plagiarism Detection...")
    print("=" * 60)
    print()
    
    # Create original assignment
    original_file = create_original_assignment()
    
    # Create plagiarized assignment
    plagiarized_file = create_plagiarized_assignment()
    
    # Create completely different assignment
    different_file = create_completely_different_assignment()
    
    print()
    print("📋 Test Files Created Successfully!")
    print("=" * 40)
    print(f"📄 {original_file}")
    print("   • Original content about online learning")
    print("   • Student: John Smith")
    print("   • Expected plagiarism score: LOW (0-20%)")
    print()
    print(f"📄 {plagiarized_file}")
    print("   • Plagiarized content (90%+ copied from original)")
    print("   • Student: Sarah Johnson")
    print("   • Expected plagiarism score: HIGH (70-95%)")
    print()
    print(f"📄 {different_file}")
    print("   • Completely different content about climate change")
    print("   • Student: Michael Brown")
    print("   • Expected plagiarism score: LOW (0-20%)")
    print()
    print("🎯 How to Test:")
    print("=" * 20)
    print("1. Go to your E-Assignment system")
    print("2. Create a new assignment")
    print("3. Upload the ORIGINAL file first")
    print("4. Then upload the PLAGIARIZED file")
    print("5. Check the plagiarism scores!")
    print("6. Upload the DIFFERENT file to see low similarity")
    print()
    print("🚀 Expected Results:")
    print("   • Original: 0-20% similarity")
    print("   • Plagiarized: 70-95% similarity")
    print("   • Different: 0-20% similarity")

if __name__ == "__main__":
    main()

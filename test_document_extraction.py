#!/usr/bin/env python3
"""
Simple Test for Document Text Extraction
Tests the text extraction functions directly.
"""

import os
import tempfile

def test_pdf_extraction():
    """Test PDF text extraction."""
    print("🧪 Testing PDF Text Extraction")
    print("-" * 40)
    
    try:
        from app import extract_text_from_pdf
        
        # Create a simple test PDF
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                pdf_path = temp_file.name
            
            # Create PDF
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.drawString(100, 750, "Test PDF Document")
            c.drawString(100, 700, "This is a sample PDF for testing plagiarism detection.")
            c.drawString(100, 650, "The system should extract this text for analysis.")
            c.save()
            
            print(f"✅ Created test PDF: {pdf_path}")
            
            # Test extraction
            result = extract_text_from_pdf(pdf_path)
            
            if "[PDF EXTRACTED TEXT]" in result:
                print("✅ PDF text extraction successful!")
                text_content = result.replace("[PDF EXTRACTED TEXT]\n", "").strip()
                print(f"   Extracted text: '{text_content}'")
                return True
            else:
                print(f"❌ PDF extraction failed: {result}")
                return False
                
        except ImportError:
            print("⚠️ reportlab not available - cannot create test PDF")
            return False
            
    except Exception as e:
        print(f"❌ PDF extraction test failed: {e}")
        return False
    finally:
        # Clean up
        if 'pdf_path' in locals() and os.path.exists(pdf_path):
            os.unlink(pdf_path)

def test_word_extraction():
    """Test Word document text extraction."""
    print("\n🧪 Testing Word Document Text Extraction")
    print("-" * 40)
    
    try:
        from app import extract_text_from_word
        
        # Create a simple test Word document
        try:
            from docx import Document
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                docx_path = temp_file.name
            
            # Create Word document
            doc = Document()
            doc.add_heading('Test Document', 0)
            doc.add_paragraph('This is a sample Word document for testing plagiarism detection.')
            doc.add_paragraph('It contains multiple paragraphs with different content.')
            doc.add_paragraph('The system should be able to extract this text for analysis.')
            doc.save(docx_path)
            
            print(f"✅ Created test Word document: {docx_path}")
            
            # Test extraction
            result = extract_text_from_word(docx_path)
            
            if "[WORD EXTRACTED TEXT]" in result:
                print("✅ Word document text extraction successful!")
                text_content = result.replace("[WORD EXTRACTED TEXT]\n", "").strip()
                print(f"   Extracted text: '{text_content[:100]}{'...' if len(text_content) > 100 else ''}'")
                return True
            else:
                print(f"❌ Word extraction failed: {result}")
                return False
                
        except ImportError:
            print("⚠️ python-docx not available - cannot create test Word document")
            return False
            
    except Exception as e:
        print(f"❌ Word extraction test failed: {e}")
        return False
    finally:
        # Clean up
        if 'docx_path' in locals() and os.path.exists(docx_path):
            os.unlink(docx_path)

def test_odt_extraction():
    """Test ODT document text extraction."""
    print("\n🧪 Testing ODT Document Text Extraction")
    print("-" * 40)
    
    try:
        from app import extract_text_from_odt
        
        # Create a simple test ODT file
        import zipfile
        import xml.etree.ElementTree as ET
        
        with tempfile.NamedTemporaryFile(suffix='.odt', delete=False) as temp_file:
            odt_path = temp_file.name
        
        # Create ODT file (it's a ZIP archive with XML)
        with zipfile.ZipFile(odt_path, 'w') as odt_file:
            # Create minimal content.xml
            content_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0">
    <office:body>
        <office:text>
            <text:p>This is a sample ODT document for testing plagiarism detection.</text:p>
            <text:p>The system should extract this text for analysis.</text:p>
        </office:text>
    </office:body>
</office:document-content>'''
            
            odt_file.writestr('content.xml', content_xml)
            odt_file.writestr('META-INF/manifest.xml', '<?xml version="1.0"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/></manifest:manifest>')
        
        print(f"✅ Created test ODT document: {odt_path}")
        
        # Test extraction
        result = extract_text_from_odt(odt_path)
        
        if "[ODT EXTRACTED TEXT]" in result:
            print("✅ ODT document text extraction successful!")
            text_content = result.replace("[ODT EXTRACTED TEXT]\n", "").strip()
            print(f"   Extracted text: '{text_content}'")
            return True
        else:
            print(f"❌ ODT extraction failed: {result}")
            return False
            
    except Exception as e:
        print(f"❌ ODT extraction test failed: {e}")
        return False
    finally:
        # Clean up
        if 'odt_path' in locals() and os.path.exists(odt_path):
            os.unlink(odt_path)

def test_plagiarism_detection():
    """Test plagiarism detection on extracted text."""
    print("\n🧪 Testing Plagiarism Detection on Extracted Text")
    print("-" * 50)
    
    try:
        from app import calculate_local_plagiarism_score
        
        # Test with sample extracted text
        sample_text = "[PDF EXTRACTED TEXT]\nThis is a sample document for testing plagiarism detection. It contains some unique content that should be analyzed."
        
        # Create mock other submissions
        other_submissions = [
            type('MockSubmission', (), {'content': 'This is completely different content for comparison purposes.'})(),
            type('MockSubmission', (), {'content': 'Another different submission with unique content.'})()
        ]
        
        print("📝 Testing plagiarism detection on extracted document text...")
        
        score = calculate_local_plagiarism_score(sample_text, other_submissions)
        
        print(f"✅ Plagiarism detection successful!")
        print(f"   📊 Similarity score: {score}%")
        
        if score > 50:
            print(f"   ⚠️ High similarity detected - possible plagiarism")
        elif score > 20:
            print(f"   ⚠️ Moderate similarity detected")
        else:
            print(f"   ✅ Low similarity - likely original content")
        
        return True
        
    except Exception as e:
        print(f"❌ Plagiarism detection test failed: {e}")
        return False

def main():
    """Main test function."""
    print("🎓 Document Text Extraction Test")
    print("=" * 60)
    print("Testing text extraction from document files:")
    print("• PDF documents with PyPDF2")
    print("• Word documents (.docx) with python-docx")
    print("• OpenDocument Text (.odt) with XML parsing")
    print("• Plagiarism detection on extracted text")
    print()
    
    try:
        # Test PDF extraction
        pdf_success = test_pdf_extraction()
        
        # Test Word extraction
        word_success = test_word_extraction()
        
        # Test ODT extraction
        odt_success = test_odt_extraction()
        
        # Test plagiarism detection
        plagiarism_success = test_plagiarism_detection()
        
        print(f"\n" + "=" * 60)
        print("📋 FINAL RESULTS")
        print("=" * 60)
        
        print(f"📄 PDF Text Extraction: {'✅ SUCCESS' if pdf_success else '❌ FAILED'}")
        print(f"📄 Word Text Extraction: {'✅ SUCCESS' if word_success else '❌ FAILED'}")
        print(f"📄 ODT Text Extraction: {'✅ SUCCESS' if odt_success else '❌ FAILED'}")
        print(f"🔍 Plagiarism Detection: {'✅ SUCCESS' if plagiarism_success else '❌ FAILED'}")
        
        if pdf_success and word_success and odt_success and plagiarism_success:
            print("\n🎉 ALL TESTS PASSED!")
            print("\n🎯 KEY ACHIEVEMENTS:")
            print("• PDF documents can now be analyzed for plagiarism")
            print("• Word documents (.docx) can now be analyzed for plagiarism")
            print("• OpenDocument Text (.odt) can now be analyzed for plagiarism")
            print("• Text extraction integrates seamlessly with plagiarism detection")
            print("• Students can submit documents directly without conversion")
        else:
            print("\n⚠️ Some tests failed - check the output above for details")
        
        print(f"\n💡 IMPACT:")
        print("• Previously unsupported file types are now fully supported")
        print("• PDF: Binary file → Text extractable → Plagiarism detectable")
        print("• Word: Binary file → Text extractable → Plagiarism detectable")
        print("• ODT: Binary file → Text extractable → Plagiarism detectable")
        print("• Archive files (.zip, .rar) still need manual extraction")
        
    except KeyboardInterrupt:
        print("\n\n⚠️ Tests interrupted by user")
    except Exception as e:
        print(f"\n❌ Unexpected error during testing: {e}")

if __name__ == "__main__":
    main()

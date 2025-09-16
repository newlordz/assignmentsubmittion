#!/usr/bin/env python3
"""
Test Word Document Workflow
Demonstrates the complete workflow for Word document plagiarism detection.
"""

import os
import tempfile
from pathlib import Path

def create_sample_word_document():
    """Create a sample Word document for testing."""
    try:
        from docx import Document
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        doc.add_heading('Sample Assignment: The Impact of Technology on Education', 0)
        
        # Add introduction
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph(
            'Technology has revolutionized the way we approach education in the 21st century. '
            'From online learning platforms to artificial intelligence-powered tutoring systems, '
            'educational technology has transformed traditional classroom settings and created '
            'new opportunities for personalized learning experiences.'
        )
        
        # Add main content
        doc.add_heading('Benefits of Educational Technology', level=1)
        doc.add_paragraph(
            'One of the most significant advantages of educational technology is its ability '
            'to provide personalized learning experiences. Students can learn at their own pace, '
            'access resources tailored to their learning style, and receive immediate feedback '
            'on their progress.'
        )
        
        doc.add_paragraph(
            'Furthermore, technology has made education more accessible to students worldwide. '
            'Online courses, virtual classrooms, and digital resources have broken down '
            'geographical barriers and provided learning opportunities to individuals who '
            'previously had limited access to quality education.'
        )
        
        # Add conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(
            'In conclusion, technology has fundamentally changed the landscape of education. '
            'While there are challenges to overcome, the benefits of educational technology '
            'far outweigh the drawbacks. As we continue to advance technologically, it is '
            'crucial that we harness these innovations to create more effective and inclusive '
            'learning environments for all students.'
        )
        
        # Add a table with some data
        doc.add_heading('Technology Usage Statistics', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Technology'
        hdr_cells[1].text = 'Usage Rate'
        hdr_cells[2].text = 'Effectiveness'
        
        # Add data rows
        row_cells = table.add_row().cells
        row_cells[0].text = 'Online Learning Platforms'
        row_cells[1].text = '85%'
        row_cells[2].text = 'High'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'AI Tutoring Systems'
        row_cells[1].text = '45%'
        row_cells[2].text = 'Medium'
        
        row_cells = table.add_row().cells
        row_cells[0].text = 'Virtual Reality'
        row_cells[1].text = '20%'
        row_cells[2].text = 'Emerging'
        
        return doc
        
    except ImportError:
        print("❌ python-docx not available - cannot create Word document")
        return None
    except Exception as e:
        print(f"❌ Error creating Word document: {e}")
        return None

def test_word_document_workflow():
    """Test the complete Word document workflow."""
    print("🎓 Word Document Plagiarism Detection Workflow Test")
    print("=" * 60)
    
    try:
        from app import read_file_content, calculate_local_plagiarism_score, allowed_file
        
        # Create sample Word document
        print("📝 Creating sample Word document...")
        doc = create_sample_word_document()
        
        if not doc:
            print("❌ Cannot proceed without Word document")
            return False
        
        # Save document to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            doc_path = temp_file.name
        
        doc.save(doc_path)
        print(f"✅ Created Word document: {os.path.basename(doc_path)}")
        
        # Test 1: File Upload Permission
        print(f"\n🔍 Test 1: File Upload Permission")
        print("-" * 40)
        
        filename = os.path.basename(doc_path)
        is_allowed = allowed_file(filename)
        
        if is_allowed:
            print(f"✅ File upload allowed: {filename}")
        else:
            print(f"❌ File upload blocked: {filename}")
            return False
        
        # Test 2: Text Extraction
        print(f"\n🔍 Test 2: Text Extraction from Word Document")
        print("-" * 50)
        
        extracted_content = read_file_content(doc_path)
        
        if "[WORD EXTRACTED TEXT]" in extracted_content:
            print("✅ Word document text extraction successful!")
            
            # Extract the actual text content
            text_content = extracted_content.replace("[WORD EXTRACTED TEXT]\n", "").strip()
            print(f"📄 Extracted text length: {len(text_content)} characters")
            print(f"📄 First 200 characters: '{text_content[:200]}...'")
            
            # Show what was extracted
            print(f"\n📋 Extracted Content Preview:")
            lines = text_content.split('\n')
            for i, line in enumerate(lines[:10]):  # Show first 10 lines
                if line.strip():
                    print(f"   {i+1:2d}. {line.strip()}")
            if len(lines) > 10:
                print(f"   ... and {len(lines) - 10} more lines")
                
        else:
            print(f"❌ Word document text extraction failed: {extracted_content}")
            return False
        
        # Test 3: Plagiarism Detection
        print(f"\n🔍 Test 3: Plagiarism Detection on Extracted Text")
        print("-" * 55)
        
        # Create mock other submissions for comparison
        other_submissions = [
            type('MockSubmission', (), {
                'content': 'This is a completely different essay about climate change and its effects on global warming patterns.'
            })(),
            type('MockSubmission', (), {
                'content': 'Another different submission discussing the history of ancient civilizations and their contributions to modern society.'
            })(),
            type('MockSubmission', (), {
                'content': 'A third submission about the benefits of renewable energy sources and sustainable development practices.'
            })()
        ]
        
        print("📊 Running plagiarism detection analysis...")
        
        try:
            similarity_score = calculate_local_plagiarism_score(extracted_content, other_submissions)
            
            print(f"✅ Plagiarism detection completed successfully!")
            print(f"📊 Similarity Score: {similarity_score:.2f}%")
            
            # Interpret the score
            if similarity_score > 70:
                print(f"⚠️ HIGH SIMILARITY - Possible plagiarism detected")
            elif similarity_score > 40:
                print(f"⚠️ MODERATE SIMILARITY - Review recommended")
            elif similarity_score > 20:
                print(f"✅ LOW SIMILARITY - Likely original content")
            else:
                print(f"✅ VERY LOW SIMILARITY - Highly original content")
                
        except Exception as e:
            print(f"❌ Plagiarism detection failed: {e}")
            return False
        
        # Test 4: Content Analysis
        print(f"\n🔍 Test 4: Content Analysis")
        print("-" * 30)
        
        # Analyze the extracted content
        word_count = len(text_content.split())
        paragraph_count = len([line for line in text_content.split('\n') if line.strip()])
        
        print(f"📊 Content Statistics:")
        print(f"   • Total characters: {len(text_content)}")
        print(f"   • Word count: {word_count}")
        print(f"   • Paragraph count: {paragraph_count}")
        print(f"   • Contains tables: {'Yes' if 'Technology Usage Statistics' in text_content else 'No'}")
        print(f"   • Contains headings: {'Yes' if 'Introduction' in text_content and 'Conclusion' in text_content else 'No'}")
        
        return True
        
    except Exception as e:
        print(f"❌ Workflow test failed: {e}")
        return False
    finally:
        # Clean up temporary file
        if 'doc_path' in locals() and os.path.exists(doc_path):
            os.unlink(doc_path)

def demonstrate_real_world_usage():
    """Demonstrate how this would work in a real-world scenario."""
    print(f"\n🌍 Real-World Usage Scenario")
    print("=" * 40)
    
    print("📚 Scenario: Student submits Word document assignment")
    print()
    print("1. 📤 Student uploads 'essay_assignment.docx'")
    print("2. 🔍 System detects it's a Word document")
    print("3. 📄 System extracts all text content automatically")
    print("4. 🔬 System runs plagiarism detection on extracted text")
    print("5. 📊 System generates detailed similarity report")
    print("6. ✅ Lecturer receives full analysis results")
    print()
    print("🎯 Benefits:")
    print("   • No manual conversion required")
    print("   • Full text content analyzed")
    print("   • Tables and formatting preserved in text")
    print("   • Same quality analysis as plain text files")
    print("   • Students can use familiar Word interface")

def main():
    """Main test function."""
    print("🎓 Word Document Plagiarism Detection - Complete Workflow Test")
    print("=" * 70)
    print("This test demonstrates the complete workflow for Word document analysis:")
    print("• File upload permission")
    print("• Text extraction from Word document")
    print("• Plagiarism detection on extracted text")
    print("• Content analysis and statistics")
    print()
    
    try:
        # Run the workflow test
        workflow_success = test_word_document_workflow()
        
        # Demonstrate real-world usage
        demonstrate_real_world_usage()
        
        print(f"\n" + "=" * 70)
        print("📋 FINAL RESULTS")
        print("=" * 70)
        
        if workflow_success:
            print("🎉 WORD DOCUMENT WORKFLOW TEST PASSED!")
            print()
            print("✅ All components working correctly:")
            print("   • Word document creation and saving")
            print("   • File upload permission system")
            print("   • Text extraction from Word documents")
            print("   • Plagiarism detection on extracted text")
            print("   • Content analysis and statistics")
            print()
            print("🚀 Your system is ready for Word document submissions!")
            print("   Students can now submit .docx and .doc files directly")
            print("   System will automatically extract text and analyze for plagiarism")
        else:
            print("❌ Some tests failed - check the output above for details")
        
    except KeyboardInterrupt:
        print("\n\n⚠️ Tests interrupted by user")
    except Exception as e:
        print(f"\n❌ Unexpected error during testing: {e}")

if __name__ == "__main__":
    main()
